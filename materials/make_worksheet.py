# -*- coding: utf-8 -*-
"""
직업계고 1학년 「국어 × 디자인」 읽기 독해 수업용 인쇄 A4 활동지 생성 스크립트.

명세(SSOT): 02_활동지.md
방법론: .claude/skills/docx-worksheet/SKILL.md (set_korean_font / 머리말 음영 / 손글씨 표)
공통 맥락: _workspace/00_shared-brief.md (쉬운 한국어·격려 톤, AI 윤리 3원칙)

실행:  python materials\\make_worksheet.py
출력:  materials/output/활동지_디자인직무읽기.docx
"""
import os

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT

FONT = "맑은 고딕"


# ---------------------------------------------------------------------------
# 공통 헬퍼 (스킬 패턴)
# ---------------------------------------------------------------------------
def set_korean_font(run, name=FONT, size=11, bold=False, color=None):
    """라틴·East Asian 양쪽에 한글 폰트를 지정해 글리프 깨짐을 막는다."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)
    rfonts.set(qn('w:eastAsia'), name)


def shade_cell(cell, fill="EFEFEF"):
    """셀에 옅은 음영(인쇄·복사 잘 되게)."""
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(sh)


def set_row_height(row, mm, rule=WD_ROW_HEIGHT_RULE.AT_LEAST):
    row.height_rule = rule
    row.height = Mm(mm)


def cell_text(cell, text, size=11, bold=False, color=None, align=None,
              space_after=0):
    """셀의 첫 문단을 비우고 텍스트를 채운다(한글 폰트 적용)."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_korean_font(r, size=size, bold=bold, color=color)
    return p


def add_para(doc, text="", size=11, bold=False, color=None, align=None,
             space_before=2, space_after=2):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_korean_font(r, size=size, bold=bold, color=color)
    return p


GUIDE_COLOR = RGBColor(0x55, 0x55, 0x55)  # 안내 문구용 회색


def add_guide(doc, text):
    """항목 아래에 붙는 짧고 친절한 안내 문구(회색, 작은 글씨)."""
    return add_para(doc, "→ " + text, size=9, color=GUIDE_COLOR,
                    space_before=1, space_after=3)


def section_heading(doc, text):
    """번호 붙은 구역 제목(옅은 회색 음영 막대)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    set_table_full_width(t)
    set_row_height(t.rows[0], 8)
    c = t.rows[0].cells[0]
    shade_cell(c, "DDE3EA")
    cell_text(c, text, size=12, bold=True)
    add_para(doc, "", size=2, space_before=0, space_after=0)
    return t


def set_table_full_width(t):
    t.allow_autofit = False
    tblPr = t._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')  # 100%
    tblPr.append(tblW)


def set_col_widths(t, widths_mm):
    for row in t.rows:
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)


def add_box(doc, height_mm, fill=None):
    """손으로 채우는 빈 박스(1x1 표)."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    set_table_full_width(t)
    set_row_height(t.rows[0], height_mm)
    if fill:
        shade_cell(t.rows[0].cells[0], fill)
    return t


def header_row_cells(t, headers):
    """표의 첫 행을 헤더로(음영+굵게)."""
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade_cell(c, "EFEFEF")
        cell_text(c, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_write_table(doc, headers, rows=4, row_h_mm=14, widths_mm=None,
                    numbered=False):
    """손글씨용 표: 헤더 + 빈 행(행 높이 넉넉히)."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_full_width(t)
    header_row_cells(t, headers)
    set_row_height(t.rows[0], 8)
    for n in range(rows):
        row = t.add_row()
        set_row_height(row, row_h_mm)
        if numbered:
            cell_text(row.cells[0], str(n + 1),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    if widths_mm:
        set_col_widths(t, widths_mm)
    return t


# ---------------------------------------------------------------------------
# 문서 빌드
# ---------------------------------------------------------------------------
def build():
    doc = Document()

    # 기본 스타일에 한글 폰트
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    style.element.rPr.rFonts.set(qn('w:ascii'), FONT)
    style.element.rPr.rFonts.set(qn('w:hAnsi'), FONT)

    # A4 세로 + 여백
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(14)
    sec.left_margin = sec.right_margin = Mm(16)

    # -------------------------------------------------------------------
    # 머리말 (옅은 음영 박스): 제목 / 학습목표 / 학번·이름·날짜
    # -------------------------------------------------------------------
    head = doc.add_table(rows=3, cols=1)
    head.style = "Table Grid"
    set_table_full_width(head)

    # 제목 행
    set_row_height(head.rows[0], 12)
    shade_cell(head.rows[0].cells[0], "E3E8EF")
    cell_text(head.rows[0].cells[0], "디자인 직무 글, 제대로 읽기",
              size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 학습목표 행
    set_row_height(head.rows[1], 10)
    shade_cell(head.rows[1].cells[0], "F2F4F7")
    cell_text(
        head.rows[1].cells[0],
        "학습목표 : 디자인 직무 글을 끝까지 읽고, 핵심 정보와 글쓴이의 의도를 "
        "내 말로 정리할 수 있다. (AI는 보조 도구!)",
        size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 학번/이름/날짜 행 (셀 안에 한 줄)
    set_row_height(head.rows[2], 10)
    cell_text(head.rows[2].cells[0],
              "  1학년      반      번   ｜  이름 :                    "
              "｜  날짜 : 2026.        .        .",
              size=11)

    add_para(doc, "", size=4, space_before=0, space_after=2)

    # -------------------------------------------------------------------
    # 0. 오늘 읽을 글
    # -------------------------------------------------------------------
    section_heading(doc, "0. 오늘 읽을 글")
    add_guide(doc, "선생님이 나눠 준 읽기 자료를 아래 칸에 붙이거나, "
                   "‘별지 참고’에 표시하고 천천히 펼쳐 보세요.")
    glue = doc.add_table(rows=1, cols=1)
    glue.style = "Table Grid"
    set_table_full_width(glue)
    set_row_height(glue.rows[0], 26)
    cell_text(glue.rows[0].cells[0],
              "[ 읽기 자료를 여기에 붙이세요 ]      □ 별지 참고",
              size=10, color=GUIDE_COLOR)

    add_para(doc, "", size=4)

    # -------------------------------------------------------------------
    # 1. 스스로 읽기
    # -------------------------------------------------------------------
    section_heading(doc, "1. 스스로 읽기 — 먼저 나 혼자 읽어요")
    add_guide(doc, "AI 도움 없이 끝까지 읽어요. 모르는 낱말·문장이 나오면 "
                   "표시하고, 뜻을 ‘짐작’해서 적어 보세요. 틀려도 괜찮아요!")
    t1 = add_write_table(
        doc,
        ["모르는 낱말 · 문장", "내가 짐작한 뜻 (틀려도 OK)"],
        rows=4, row_h_mm=14, widths_mm=[70, 108])

    add_para(doc, "", size=4)

    # -------------------------------------------------------------------
    # 2. AI 독해 도우미 활용
    # -------------------------------------------------------------------
    section_heading(doc, "2. AI 독해 도우미 활용 — 막힌 부분에 도움받기")
    add_guide(doc, "AI 독해 도우미 웹앱으로 어휘·요약·핵심어를 확인해요. "
                   "개인정보(이름·연락처)는 입력하지 않아요.")

    # 2-(1) 어휘 뜻 정리 표 (3열)
    add_para(doc, "(1) 어휘 뜻 정리", size=11, bold=True,
             space_before=4, space_after=1)
    add_guide(doc, "AI가 알려준 쉬운 뜻을 보고, 마지막 칸에는 꼭 ‘내 문장’으로 "
                   "바꿔 써요. 그래야 진짜 내 것이 돼요.")
    t2 = add_write_table(
        doc,
        ["낱말", "AI가 알려준 쉬운 뜻", "내 문장으로 다시 쓰기"],
        rows=4, row_h_mm=14, widths_mm=[34, 72, 72])

    # 2-(2) 문단 요약 표 + "베끼지 말고 내 말로" 안내
    add_para(doc, "(2) 문단 요약", size=11, bold=True,
             space_before=5, space_after=1)
    add_guide(doc, "AI 요약을 그대로 베끼지 말고, 자기 말로 한 번 더 정리해요. "
                   "AI는 틀릴 수 있으니 원문과 꼭 대조해요!")
    t3 = add_write_table(
        doc,
        ["문단", "핵심 내용을 ‘내 말로’ 정리"],
        rows=4, row_h_mm=15, widths_mm=[26, 152], numbered=True)

    # 2-(3) 핵심어
    add_para(doc, "(3) 핵심어 3~5개 적기", size=11, bold=True,
             space_before=5, space_after=1)
    add_guide(doc, "이 글에서 가장 중요한 낱말을 3~5개 골라 적어요.")
    kw = doc.add_table(rows=1, cols=5)
    kw.style = "Table Grid"
    set_table_full_width(kw)
    set_row_height(kw.rows[0], 14)
    for i in range(5):
        cell_text(kw.rows[0].cells[i], f"{i+1}.", size=10, color=GUIDE_COLOR)

    add_para(doc, "", size=4)

    # -------------------------------------------------------------------
    # 3. 핵심 정리
    # -------------------------------------------------------------------
    section_heading(doc, "3. 핵심 정리 — 글의 알맹이를 잡아요")
    add_guide(doc, "앞에서 읽고 정리한 내용을 모아, 아래 세 가지로 정리해 보세요.")
    t4 = doc.add_table(rows=3, cols=2)
    t4.style = "Table Grid"
    set_table_full_width(t4)
    labels = [
        "① 이 글의 중심 내용",
        "② 글쓴이가 말하려는 것 (의도)",
        "③ 이 글이 디자인 직무에 어떻게 쓰일까?",
    ]
    for i, lab in enumerate(labels):
        set_row_height(t4.rows[i], 18)
        shade_cell(t4.rows[i].cells[0], "F2F4F7")
        cell_text(t4.rows[i].cells[0], lab, size=10.5, bold=True)
        t4.rows[i].cells[1].text = ""
    set_col_widths(t4, [56, 122])

    add_para(doc, "", size=4)

    # -------------------------------------------------------------------
    # 4. 이해 점검
    # -------------------------------------------------------------------
    section_heading(doc, "4. 이해 점검 — 질문에 답하며 확인해요")
    add_guide(doc, "AI 도우미가 낸 점검 질문을 왼쪽에 적고, 오른쪽에 내 답을 "
                   "써요. 답이 막히면 원문을 다시 읽어 봐요.")
    t5 = add_write_table(
        doc,
        ["AI가 낸 질문", "내 답"],
        rows=3, row_h_mm=18, widths_mm=[78, 100], numbered=True)

    add_para(doc, "", size=4)

    # -------------------------------------------------------------------
    # 5. 한 문장 요약
    # -------------------------------------------------------------------
    section_heading(doc, "5. 한 문장 요약 — 오늘 읽은 글을 한 문장으로")
    add_guide(doc, "오늘 읽은 글 전체를 딱 한 문장으로 줄여 보세요. "
                   "‘이 글은 ~ 에 대한 글이다.’ 처럼 써도 좋아요.")
    add_box(doc, 18)

    add_para(doc, "", size=4)

    # -------------------------------------------------------------------
    # 6. 성찰 (SEL / AI 윤리 연계)
    # -------------------------------------------------------------------
    section_heading(doc, "6. 성찰 — AI 도우미, 잘 썼나요?")
    add_guide(doc, "AI는 ‘보조 도구’예요. 틀릴 수도 있어요. 오늘 내가 AI를 어떻게 "
                   "썼는지 솔직하게 돌아봐요.")
    t6 = doc.add_table(rows=3, cols=2)
    t6.style = "Table Grid"
    set_table_full_width(t6)
    reflect = [
        "AI 도우미가 도움이 된 점",
        "아쉽거나 틀린 점 (원문과 다른 점)",
        "그래서 나는 AI를 어떻게 쓰면 좋을까?",
    ]
    for i, lab in enumerate(reflect):
        set_row_height(t6.rows[i], 18)
        shade_cell(t6.rows[i].cells[0], "F2F4F7")
        cell_text(t6.rows[i].cells[0], lab, size=10.5, bold=True)
        t6.rows[i].cells[1].text = ""
    set_col_widths(t6, [60, 118])

    # 하단 AI 윤리 3원칙 안내 박스
    add_para(doc, "", size=3)
    tip = doc.add_table(rows=1, cols=1)
    tip.style = "Table Grid"
    set_table_full_width(tip)
    set_row_height(tip.rows[0], 12)
    shade_cell(tip.rows[0].cells[0], "FFF6E5")
    cell_text(
        tip.rows[0].cells[0],
        "기억해요!  ① 이름·연락처 같은 개인정보는 입력하지 않기   "
        "② AI는 틀릴 수 있으니 원문과 대조하기   "
        "③ 베끼지 말고 ‘내 말로’ 정리하기",
        size=9.5, bold=True)

    # -------------------------------------------------------------------
    # 저장
    # -------------------------------------------------------------------
    out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "활동지_디자인직무읽기.docx")
    doc.save(out_path)
    print("저장 완료:", out_path)
    return out_path


if __name__ == "__main__":
    build()
